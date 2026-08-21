#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Научный монитор для Йоды: поиск статей (Europe PMC) и ПОЛНЫЙ подробный
Word-разбор статьи по запросу через excash gemini-3.1-pro (резерв — DeepSeek V4 Pro).

Полнотекст: Europe PMC (OA) → OpenAlex + Unpaywall (OA PDF/HTML по DOI) → извлечение
(pdftotext / BeautifulSoup) — методика позаимствована из DocMed (svoivra4.ru/evidence).

  sci.py search "<запрос>" [--n 8]
  sci.py word "<запрос или PMID/DOI>" [--pmid 123] [--doi 10.x/y] [--send]
"""
import argparse, html, json, os, re, subprocess, sys, tempfile, time
from datetime import datetime, timedelta, timezone
from io import BytesIO
from urllib.parse import quote
import requests

HERE = os.path.dirname(os.path.abspath(__file__))
TOME = os.path.expanduser("~/.openclaw/workspace/skills/tome/tome.py")
PYBIN = os.path.expanduser("~/mailvenv/bin/python")
EPMC = "https://www.ebi.ac.uk/europepmc/webservices/rest"
OPENALEX = "https://api.openalex.org/works"
UNPAYWALL = "https://api.unpaywall.org/v2"
FT_MIN, FT_MAX, FT_MAX_BYTES = 1500, 200000, 25 * 1024 * 1024

ENV = {}
_p = os.path.expanduser("~/.openclaw/.env")
if os.path.exists(_p):
    for line in open(_p, encoding="utf-8"):
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            k, v = line.split("=", 1)
            ENV[k.strip()] = v.strip().strip('"').strip("'")

UNPAYWALL_EMAIL = ENV.get("UNPAYWALL_EMAIL", "you@example.com")

# ---------- LLM: excash gemini-3.1-pro -> резерв DeepSeek V4 Pro ----------
def _llm_once(base_url, key, model, messages, max_tokens, temperature, timeout=600):
    r = requests.post(base_url.rstrip("/") + "/chat/completions",
                      headers={"Authorization": "Bearer " + key, "Content-Type": "application/json"},
                      json={"model": model, "messages": messages,
                            "max_tokens": max_tokens, "temperature": temperature}, timeout=timeout)
    r.raise_for_status()
    return ((r.json().get("choices", [{}])[0].get("message", {}) or {}).get("content", "")) or ""

def llm(messages, max_tokens=20000, temperature=0.35):
    key, url = ENV.get("EXCASH_API_KEY"), ENV.get("EXCASH_API_URL")
    if key and url:
        try:
            txt = _llm_once(url, key, "gemini-3.1-pro", messages, max_tokens, temperature)
            if txt.strip():
                return txt
        except Exception as e:
            sys.stderr.write(f"excash недоступен -> DeepSeek V4 Pro: {e}\n")
    dk = ENV.get("DEEPSEEK_API_KEY")
    if dk:
        # deepseek-v4-pro — reasoning-модель: щедрый лимит, чтобы content заполнился после рассуждений
        return _llm_once("https://api.deepseek.com/v1", dk, "deepseek-v4-pro",
                         messages, min(max(max_tokens, 8000), 32000), temperature, timeout=900)
    sys.exit("нет ни EXCASH, ни DEEPSEEK ключей в ~/.openclaw/.env")

# ---------- Europe PMC ----------
def epmc_search(query, n=8):
    r = requests.get(f"{EPMC}/search", params={
        "query": query, "format": "json", "resultType": "core",
        "pageSize": str(n), "sort": "P_PDATE_D desc"}, timeout=40)
    r.raise_for_status()
    return r.json().get("resultList", {}).get("result", [])

def epmc_one(spec, pmid=None, doi=None):
    if pmid:
        q = f"EXT_ID:{pmid} AND SRC:MED"
    elif doi:
        q = f'DOI:"{doi}"'
    elif re.fullmatch(r"\d{6,9}", spec or ""):
        q = f"EXT_ID:{spec} AND SRC:MED"
    elif spec and spec.upper().startswith("PMC"):
        q = f"PMCID:{spec.upper()}"
    elif spec and "/" in spec and spec[0].isdigit():
        q = f'DOI:"{spec}"'
    else:
        q = spec
    res = epmc_search(q, n=1)
    return res[0] if res else None

def fmt_meta(a):
    j = a.get("journalInfo", {}).get("journal", {}) or {}
    return {
        "title": a.get("title", "").rstrip("."),
        "authors": a.get("authorString", ""),
        "journal": j.get("title", ""),
        "year": a.get("pubYear", ""),
        "doi": a.get("doi", ""),
        "pmid": a.get("pmid", ""),
        "pmcid": a.get("pmcid", ""),
        "oa": a.get("isOpenAccess") == "Y",
        "abstract": re.sub(r"(?s)<[^>]+>", " ", a.get("abstractText", "") or ""),
    }

# ---------- Полнотекст (методика DocMed) ----------
def _pmc_xml_text(pmcid):
    try:
        r = requests.get(f"{EPMC}/PMC/{pmcid}/fullTextXML", timeout=60)
        if r.status_code != 200:
            return ""
        xml = re.sub(r"(?is)<(ref-list|back|table-wrap|fig|graphic).*?</\1>", " ", r.text)
        txt = html.unescape(re.sub(r"(?s)<[^>]+>", " ", xml))
        return re.sub(r"\s+", " ", txt).strip()[:FT_MAX]
    except Exception:
        return ""

def _oa_candidate_urls(doi, pmcid):
    urls = []
    if pmcid:
        urls.append(f"https://pmc.ncbi.nlm.nih.gov/articles/{pmcid}/pdf/")
    if doi:
        d = doi.lower()
        try:
            r = requests.get(f"{OPENALEX}/https://doi.org/{quote(d, safe='/:;()[]')}",
                             params={"mailto": UNPAYWALL_EMAIL}, timeout=30)
            if r.status_code == 200:
                w = r.json()
                for loc in [w.get("best_oa_location"), w.get("primary_location")] + (w.get("locations") or []):
                    if isinstance(loc, dict):
                        for u in (loc.get("pdf_url"), loc.get("landing_page_url")):
                            if u:
                                urls.append(u)
        except Exception:
            pass
        try:
            r = requests.get(f"{UNPAYWALL}/{quote(d, safe='')}", params={"email": UNPAYWALL_EMAIL}, timeout=30)
            if r.status_code == 200:
                data = r.json()
                for loc in [data.get("best_oa_location")] + (data.get("oa_locations") or []):
                    if isinstance(loc, dict):
                        for u in (loc.get("url_for_pdf"), loc.get("url_for_landing_page"), loc.get("url")):
                            if u:
                                urls.append(u)
        except Exception:
            pass
        # Semantic Scholar — ещё один легальный источник OA (часто находит авторские рукописи в репозиториях)
        try:
            r = requests.get(f"https://api.semanticscholar.org/graph/v1/paper/DOI:{quote(d, safe='')}",
                             params={"fields": "openAccessPdf"}, timeout=30)
            if r.status_code == 200:
                oa = (r.json() or {}).get("openAccessPdf") or {}
                if oa.get("url"):
                    urls.append(oa["url"])
        except Exception:
            pass
    # выкидываем pubmed/doi-редиректоры (это абстракт, не полный текст) и любые paywall-bypass
    # зеркала (как это делает сам DocMed) — работаем только с легальными OA-источниками
    SKIP = ("pubmed.ncbi.nlm.nih.gov", "://doi.org/", "ncbi.nlm.nih.gov/pubmed", "sci-hub", "scihub")
    seen, out = set(), []
    for u in urls:
        if u and u not in seen and not any(s in u.lower() for s in SKIP):
            seen.add(u); out.append(u)
    out.sort(key=lambda u: 0 if ("pdf" in u.lower()) else 1)  # PDF первыми
    return out[:12]

def _pdf_to_text(pdf_bytes):
    with tempfile.TemporaryDirectory(prefix="sci_pdf_") as tmp:
        p, t = os.path.join(tmp, "a.pdf"), os.path.join(tmp, "a.txt")
        with open(p, "wb") as f:
            f.write(pdf_bytes)
        try:
            r = subprocess.run(["pdftotext", "-layout", "-enc", "UTF-8", p, t], capture_output=True, timeout=75)
            if r.returncode == 0 and os.path.exists(t):
                txt = open(t, encoding="utf-8", errors="ignore").read()
                if len(txt.strip()) >= FT_MIN:
                    return re.sub(r"[ \t]+", " ", txt)[:FT_MAX]
        except Exception:
            pass
    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(pdf_bytes))
        parts = [(pg.extract_text() or "") for pg in reader.pages[:80]]
        return re.sub(r"\s+", " ", "\n".join(parts)).strip()[:FT_MAX]
    except Exception:
        return ""

def _html_to_text(html_text):
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_text, "html.parser")
        for n in soup(["script", "style", "noscript", "svg", "nav", "footer", "header", "form"]):
            n.decompose()
        main = (soup.find("article") or soup.find(attrs={"role": "main"}) or soup.find("main")
                or soup.find(class_=re.compile(r"(article|fulltext|content|main)", re.I)) or soup.body or soup)
        return re.sub(r"\n{3,}", "\n\n", main.get_text("\n")).strip()[:FT_MAX]
    except Exception:
        return re.sub(r"(?s)<[^>]+>", " ", html_text)[:FT_MAX]

def _fetch_url_text(url):
    """Возвращает (текст, is_pdf)."""
    try:
        r = requests.get(url, headers={
            "User-Agent": f"SvoyVrachSci/1.0 (+mailto:{UNPAYWALL_EMAIL})",
            "Accept": "application/pdf,text/html,application/xml;q=0.9,*/*;q=0.8"},
            timeout=75, allow_redirects=True)
        if r.status_code != 200:
            return "", False
        raw = r.content[:FT_MAX_BYTES]
        ct = (r.headers.get("Content-Type") or "").lower()
        final = str(r.url).lower().split("?")[0]
        if raw[:4] == b"%PDF" or "pdf" in ct or final.endswith(".pdf"):
            return _pdf_to_text(raw), True
        text = raw.decode(r.encoding or "utf-8", errors="ignore")
        if "xml" in ct or "<article" in text[:1500].lower():
            return re.sub(r"\s+", " ", html.unescape(re.sub(r"(?s)<[^>]+>", " ", text)))[:FT_MAX], False
        return _html_to_text(text), False
    except Exception:
        return "", False

def preprint_versions(title, doi=""):
    """Ищет свободную препринт-версию платной статьи: bioRxiv/medRxiv + arXiv.
    Часто именно так полный текст доступен легально."""
    urls = []
    q = (title or "")[:200]
    if not q:
        return urls
    # bioRxiv / medRxiv
    try:
        r = requests.get("https://api.biorxiv.org/covid19/0")  # прогрев/проверка доступности
    except Exception:
        pass
    for server in ("biorxiv", "medrxiv"):
        try:
            r = requests.get(f"https://api.biorxiv.org/details/{server}/{quote(doi)}", timeout=25) if doi else None
            if r is not None and r.status_code == 200:
                for it in (r.json().get("collection") or []):
                    d = it.get("doi")
                    if d:
                        urls.append(f"https://www.{server}.org/content/{d}v1.full")
        except Exception:
            pass
    # arXiv по названию
    try:
        r = requests.get("http://export.arxiv.org/api/query",
                         params={"search_query": f'ti:"{q}"', "max_results": 2}, timeout=30)
        if r.status_code == 200:
            for m in re.finditer(r"<id>(http[^<]+)</id>", r.text):
                u = m.group(1)
                if "/abs/" in u:
                    urls.append(u.replace("/abs/", "/pdf/"))
    except Exception:
        pass
    return urls[:4]


def fetch_fulltext(art):
    """Возвращает (текст, источник). Пусто если полнотекст не добыт.
    PDF принимаем от FT_MIN; HTML-лендинг — только если заметно длиннее абстракта (>=4000),
    иначе это просто страница с абстрактом, а не полный текст."""
    pmcid, doi = art.get("pmcid"), art.get("doi")
    if pmcid and art.get("isOpenAccess") == "Y":
        t = _pmc_xml_text(pmcid)
        if len(t) >= FT_MIN:
            return t, "Europe PMC (PMC XML)"
    for u in _oa_candidate_urls(doi, pmcid):
        t, is_pdf = _fetch_url_text(u)
        if t and len(t) >= (FT_MIN if is_pdf else 4000):
            return t, u
    # последняя легальная попытка — свободный препринт той же работы
    for u in preprint_versions(art.get("title", ""), doi or ""):
        t, is_pdf = _fetch_url_text(u)
        if t and len(t) >= FT_MIN:
            return t, u + "  (препринт-версия)"
    return "", ""

# ---------- docx ----------
def build_docx(title, body_md, meta, ft_src, path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(12)
    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    meta_p = doc.add_paragraph()
    meta_p.add_run(f"{meta['authors']}\n").italic = True
    meta_p.add_run(" · ".join(x for x in [meta["journal"], str(meta["year"])] if x) + "\n").bold = True
    ids = " · ".join(x for x in [f"PMID {meta['pmid']}" if meta["pmid"] else "",
                                 f"DOI {meta['doi']}" if meta["doi"] else ""] if x)
    meta_p.add_run(ids)
    src_line = ("полный текст: " + ft_src) if ft_src else ("Open Access" if meta["oa"] else "разбор по абстракту")
    meta_p.add_run("\n" + src_line).font.size = Pt(9)
    doc.add_paragraph("─" * 40)

    def add_bold(p, text):
        for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
            run = p.add_run(chunk)
            if i % 2 == 1:
                run.bold = True

    for raw in body_md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if re.match(r"^#{1,3}\s", line):
            doc.add_heading(line.lstrip("# ").strip(), level=min(len(line) - len(line.lstrip("#")), 3))
        elif re.match(r"^[-*•]\s", line):
            add_bold(doc.add_paragraph(style="List Bullet"), re.sub(r"^[-*•]\s+", "", line))
        elif re.match(r"^\d+[.)]\s", line):
            add_bold(doc.add_paragraph(style="List Number"), re.sub(r"^\d+[.)]\s+", "", line))
        else:
            add_bold(doc.add_paragraph(), line)
    doc.save(path)

SYS_PROMPT = """Ты — научный редактор для практикующего врача травматолога-ортопеда (детская ортопедия).
Составь ПОЛНЫЙ, подробный разбор научной статьи на русском языке — рабочий документ для врача,
а не короткий дайджест. Будь конкретным: приводи реальные числа, размеры выборок, значения p,
доверительные интервалы, дозы, сроки. Не выдумывай данные; если чего-то нет — пиши «в статье не указано».

Структура (Markdown-заголовки ## и списки):
## Название и выходные данные
## О чём статья коротко
## Актуальность и цель
## Дизайн и методология
## Результаты
## Выводы авторов
## Ограничения исследования
## Уровень доказательности и критическая оценка
## Клиническая значимость для практики
Пиши профессионально, плотно, без воды."""

def make_word_from_pdf(pdf_path, send=False):
    """Разбор по PDF, который дал врач (его законный доступ: подписка/библиотека/репозиторий)."""
    if not os.path.exists(pdf_path):
        sys.exit(f"файл не найден: {pdf_path}")
    with open(pdf_path, "rb") as f:
        text = _pdf_to_text(f.read())
    if len(text) < 500:
        sys.exit("не удалось извлечь текст из PDF (возможно, скан без OCR)")
    first = next((l.strip() for l in text.splitlines() if len(l.strip()) > 25), os.path.basename(pdf_path))
    m = {"title": first[:180], "authors": "", "journal": "", "year": "",
         "doi": "", "pmid": "", "pmcid": "", "oa": False, "abstract": ""}
    body = llm([{"role": "system", "content": SYS_PROMPT},
                {"role": "user", "content": f"ПОЛНЫЙ ТЕКСТ СТАТЬИ (PDF от врача):\n{text[:FT_MAX]}"}]).strip()
    body = re.sub(r"^```(?:markdown)?|```$", "", body, flags=re.M).strip()
    out = "/tmp/sci_pdf_%s.docx" % re.sub(r"[^\w\-]", "_", os.path.basename(pdf_path))[:30]
    build_docx(m["title"] or "Разбор статьи", body, m, "PDF, предоставленный врачом", out)
    if send:
        subprocess.run([PYBIN, TOME, "file", out, "--caption", f"📄 Разбор статьи: {m['title'][:90]}"], check=False)
        try:
            os.remove(out)
        except Exception:
            pass
        print("✓ Word-разбор по твоему PDF отправлен доктору")
    else:
        print(out)

def make_word(spec, pmid=None, doi=None, send=False):
    art = epmc_one(spec, pmid=pmid, doi=doi)
    if not art:
        sys.exit(f"статья не найдена: {spec or pmid or doi}")
    m = fmt_meta(art)
    ft, ft_src = fetch_fulltext(art)
    source = (f"НАЗВАНИЕ: {m['title']}\nАВТОРЫ: {m['authors']}\n"
              f"ЖУРНАЛ: {m['journal']} ({m['year']})\nDOI: {m['doi']}\n\nАБСТРАКТ:\n{m['abstract']}")
    if ft:
        source += f"\n\nПОЛНЫЙ ТЕКСТ (усечён):\n{ft}"
    body = llm([{"role": "system", "content": SYS_PROMPT}, {"role": "user", "content": source}]).strip()
    body = re.sub(r"^```(?:markdown)?|```$", "", body, flags=re.M).strip()
    safe = re.sub(r"[^\w\-]", "_", (m["pmid"] or m["doi"] or "article"))[:40]
    path = f"/tmp/sci_{safe}.docx"
    build_docx(m["title"][:120] or "Разбор статьи", body, m, ft_src, path)
    if send:
        subprocess.run([PYBIN, TOME, "file", path, "--caption", f"📄 Разбор статьи: {m['title'][:90]}"], check=False)
        try:
            os.remove(path)
        except Exception:
            pass
        print(f"✓ Word-разбор отправлен доктору" + (f" (полный текст: {ft_src})" if ft_src else " (по абстракту)"))
    else:
        print(path + (f"  [полный текст: {ft_src}]" if ft_src else "  [по абстракту]"))


# Отправители научных алертов (совпадение по адресу ИЛИ имени)
_ALERT_SENDERS = (
    "elsevier", "nature.com", "springer", "wiley", "lww.com", "ovid", "jbjs",
    "bmj.com", "sagepub", "tandfonline", "karger", "thieme", "scholar.google",
    "sciencedirect", "pubmed", "ncbi.nlm.nih.gov", "researchgate", "frontiersin",
    "mdpi", "oup.com", "oxfordjournals", "jamanetwork", "nejm", "cochrane",
    "journals", "alerts", "toc-alert", "e-alert", "editorialmanager",
)


def _gmail_alert_bodies(hours):
    """Письма-алерты из Gmail за N часов, с ТЕЛОМ (их мало — быстро)."""
    import email as _em
    import email.utils as _eu
    import imaplib
    from email.header import decode_header, make_header

    user = ENV.get("GMAIL_EMAIL") or ENV.get("SCIENCE_GMAIL_USER")
    pwd = ENV.get("GMAIL_APP_PASSWORD") or ENV.get("SCIENCE_GMAIL_PASSWORD")
    if not (user and pwd):
        sys.exit("нет доступа к Gmail (GMAIL_EMAIL/GMAIL_APP_PASSWORD в ~/.openclaw/.env) — "
                 "научные алерты прочитать нечем. Скажи это доктору, не выдумывай статьи.")

    def _dec(v):
        try:
            return str(make_header(decode_header(v or "")))
        except Exception:
            return v or ""

    M = imaplib.IMAP4_SSL("imap.gmail.com", 993)
    M.login(user, pwd)
    M.select("INBOX", readonly=True)
    since = (datetime.now() - timedelta(hours=hours + 24)).strftime("%d-%b-%Y")
    _, d = M.uid("search", None, f"(SINCE {since})")
    uids = d[0].split()[-300:]
    if not uids:
        M.logout()
        return []

    # 1) заголовки пачкой — отбираем только алерты
    _, data = M.uid("fetch", b",".join(uids),
                    "(BODY.PEEK[HEADER.FIELDS (FROM SUBJECT DATE)])")
    import re as _re
    cutoff = datetime.now(timezone.utc) - timedelta(hours=hours)
    picked = []
    for part in data:
        if not isinstance(part, tuple) or len(part) < 2:
            continue
        head = part[0].decode(errors="replace") if isinstance(part[0], bytes) else str(part[0])
        mu = _re.search(r"UID (\d+)", head)
        m = _em.message_from_bytes(part[1])
        frm, subj = _dec(m.get("From")), _dec(m.get("Subject"))
        try:
            dt = _eu.parsedate_to_datetime(m.get("Date"))
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            if dt < cutoff:
                continue
        except Exception:
            pass
        hay = (frm + " " + subj).lower()
        if any(k in hay for k in _ALERT_SENDERS):
            picked.append((mu.group(1) if mu else None, frm, subj))

    # 2) тела только у отобранных
    out = []
    for uid, frm, subj in picked:
        if not uid:
            continue
        _, bd = M.uid("fetch", uid.encode(), "(BODY.PEEK[])")
        if not bd or bd[0] is None:
            continue
        msg = _em.message_from_bytes(bd[0][1])
        body = ""
        for part in msg.walk() if msg.is_multipart() else [msg]:
            if part.get_content_type() in ("text/plain", "text/html"):
                try:
                    raw = part.get_payload(decode=True) or b""
                    txt = raw.decode(part.get_content_charset() or "utf-8", errors="replace")
                except Exception:
                    continue
                if part.get_content_type() == "text/html":
                    txt = _re.sub(r"<[^>]+>", " ", txt)
                body += txt + "\n"
        body = _re.sub(r"[ \t]+", " ", _re.sub(r"\n{3,}", "\n\n", body)).strip()
        out.append({"from": frm, "subject": subj, "body": body[:14000]})
    M.logout()
    return out


def _norm_title(t):
    return re.sub(r"[^a-z0-9 ]+", " ", (t or "").lower()).split()


def _same_article(want, got):
    """Одна ли это статья. Пересечение делим на БОЛЬШЕЕ множество слов: иначе
    короткий чужой заголовок, целиком вложенный в наш длинный, даст ложные 100%."""
    stop = {"the", "a", "an", "of", "in", "and", "for", "on", "with", "to", "from",
            "by", "at", "is", "are", "study", "case", "report", "review", "using",
            "its", "this", "that", "how", "why", "could", "would", "should"}
    a = set(_norm_title(want)) - stop
    b = set(_norm_title(got)) - stop
    if not a or not b:
        return False
    inter = len(a & b)
    return inter / max(len(a), len(b)) >= 0.7 and inter >= 4


def _resolve_link(title_en):
    """Ссылка по английскому названию: Europe PMC -> doi.org/pubmed.
    В письмах ссылки битые (quoted-printable, трекеры), поэтому резолвим сами —
    но отдаём результат ТОЛЬКО если название реально совпало."""
    if not title_en or len(title_en) < 12:
        return ""
    try:
        for query in (f'TITLE:"{title_en[:180]}"', title_en[:180]):
            for art in (epmc_search(query, n=3) or []):
                m = fmt_meta(art)
                if not _same_article(title_en, m.get("title", "")):
                    continue
                if m.get("doi"):
                    return "https://doi.org/" + m["doi"]
                if m.get("pmid"):
                    return f"https://pubmed.ncbi.nlm.nih.gov/{m['pmid']}/"
    except Exception:
        pass
    return ""


def cmd_alerts(a):
    """Все статьи из научных писем за сутки: ТОП-10 подробно + остальные по журналам."""
    mails = _gmail_alert_bodies(a.hours)
    if not mails:
        print(f"Научных писем за последние {a.hours} ч в Gmail не пришло. "
              f"Так и скажи доктору — не подменяй это поиском по PubMed.")
        return
    print(f"НАУКА ЗА {a.hours} Ч — писем: {len(mails)}")
    for m in mails:
        print(f"  • {m['subject'][:95]}")
    print("=" * 62)

    def _extract(mail):
        """Разбор ОДНОГО письма. Шлюз отдаёт 504 на большом запросе, поэтому
        письма идут по одному, а не одним блобом."""
        prompt = (
            "Ниже письмо-алерт научного журнала, пришедшее врачу — ДЕТСКОМУ "
            "ТРАВМАТОЛОГУ-ОРТОПЕДУ (интересы: детская травма и ортопедия, переломы, "
            "ПКС/мениск, дисплазия ТБС, сколиоз, плоскостопие, артроскопия, "
            "реабилитация, детская хирургия, ИИ в медицине).\n\n"
            "Выпиши ВСЕ статьи из письма — не выборку. Верни СТРОГО JSON-массив, "
            "без пояснений и без markdown-ограды. Элемент:\n"
            '{"ru":"название по-русски","en":"оригинальное название как в письме",'
            '"journal":"журнал","year":"год","desc":"2-3 предложения: о чём работа, '
            'какой дизайн, что нашли — строго по тому, что есть в письме",'
            '"rating":"fire|star|pin|none","why":"одна фраза: чем полезно именно '
            'детскому травматологу-ортопеду; для rating none — пустая строка"}\n\n'
            "rating: fire = меняет практику/прорыв; star = важно и близко к его теме; "
            "pin = интересно, косвенно; none = мимо интересов.\n"
            "Только заголовок без аннотации — в desc так и напиши "
            "«в письме только заголовок, без аннотации». Не выдумывай.\n\n"
            f"=== ПИСЬМО: {mail['subject']}\nОт: {mail['from']}\n{mail['body'][:11000]}")
        raw = llm([{"role": "user", "content": prompt}], max_tokens=12000, temperature=0.2)
        t = (raw or "").strip()
        if t.startswith("```"):
            t = t.split("```")[1]
            t = t[4:] if t.lower().startswith("json") else t
        try:
            return json.loads(t[t.find("["): t.rfind("]") + 1])
        except Exception as e:
            print(f"  ⚠️ письмо «{mail['subject'][:50]}» разобрать не удалось: {e}")
            return []

    arts = []
    for mail in mails:
        got = _extract(mail)
        print(f"  разобрано из «{mail['subject'][:55]}»: {len(got)} статей")
        arts += got
    if not arts:
        print("Ни одной статьи разобрать не удалось. Скажи доктору честно, "
              "не подменяй это поиском по PubMed.")
        return

    order = {"fire": 0, "star": 1, "pin": 2, "none": 3}
    arts.sort(key=lambda x: order.get((x.get("rating") or "none").lower(), 3))
    icon = {"fire": "🔥", "star": "⭐", "pin": "📌", "none": "⬜"}

    top = [x for x in arts if (x.get("rating") or "none").lower() != "none"][: a.top]
    rest = [x for x in arts if x not in top]

    print(f"\nВСЕГО СТАТЕЙ: {len(arts)}  |  релевантных: "
          f"{sum(1 for x in arts if (x.get('rating') or 'none').lower() != 'none')}\n")

    print(f"## ТОП-{len(top)} ДЛЯ ДЕТСКОГО ОРТОПЕДА-ТРАВМАТОЛОГА\n")
    for i, x in enumerate(top, 1):
        link = _resolve_link(x.get("en", ""))
        print(f"{i}. {icon.get((x.get('rating') or 'none').lower(), '⬜')} "
              f"{x.get('ru', '?')}")
        print(f"   {x.get('journal', '?')}, {x.get('year', '')} · {x.get('en', '')[:150]}")
        print(f"   {x.get('desc', '')}")
        if x.get("why"):
            print(f"   ▸ Зачем ему: {x['why']}")
        print(f"   🔗 {link if link else 'ссылку в базах найти не удалось'}")
        print()
        time.sleep(0.35)                       # вежливо к Europe PMC

    if rest:
        print(f"\n## ОСТАЛЬНЫЕ {len(rest)} — по журналам\n")
        by_j = {}
        for x in rest:
            by_j.setdefault(x.get("journal", "Прочее"), []).append(x)
        for j, items in sorted(by_j.items(), key=lambda kv: -len(kv[1])):
            print(f"### {j} ({len(items)})")
            for x in items:
                print(f"  {icon.get((x.get('rating') or 'none').lower(), '⬜')} "
                      f"{x.get('ru', '?')}")
                d = (x.get("desc") or "").strip()
                if d and not d.startswith("в письме только заголовок"):
                    print(f"     {d[:160]}")
                link = _resolve_link(x.get("en", "")) if a.links_all else ""
                if link:
                    print(f"     🔗 {link}")
                    time.sleep(0.35)
            print()

    print("=" * 62)
    print("Ссылки резолвятся через Europe PMC по названию — если статья свежая "
          "и ещё не проиндексирована, ссылки может не быть. Это НЕ повод "
          "выдумывать URL.")


def cmd_fulltext(a):
    art = epmc_one(a.query, pmid=a.pmid, doi=a.doi)
    if not art:
        sys.exit(f"статья не найдена: {a.query or a.pmid or a.doi}")
    m = fmt_meta(art)
    print(f"НАЗВАНИЕ: {m['title']}")
    print(f"ЖУРНАЛ: {m['journal']} {m['year']} | PMID {m['pmid']} | DOI {m['doi']}")
    ft, src = fetch_fulltext(art)
    if ft:
        print(f"ПОЛНЫЙ ТЕКСТ: ДА, {len(ft)} символов. Источник: {src}\n")
        print(ft if a.all else ft[:a.chars])
    else:
        print("ПОЛНЫЙ ТЕКСТ: НЕТ в открытом доступе (проверены Europe PMC, OpenAlex, Unpaywall, "
              "Semantic Scholar, препринты).")
        print("Доступен только абстракт:\n")
        print(m["abstract"][:4000])
        print("\n→ Если нужен полный текст: у доктора есть институциональный доступ (РНИМУ/"
              "библиотека). Пусть скачает PDF и пришлёт — тогда: sci.py word --pdf <файл> --send")


def main():
    ap = argparse.ArgumentParser()
    sub = ap.add_subparsers(dest="cmd", required=True)
    s = sub.add_parser("search"); s.add_argument("query"); s.add_argument("--n", type=int, default=8)
    al = sub.add_parser("alerts", help="все статьи из научных писем за сутки")
    al.add_argument("--hours", type=int, default=24)
    al.add_argument("--top", type=int, default=10, help="сколько статей разобрать подробно")
    al.add_argument("--links-all", action="store_true",
                    help="резолвить ссылки и для остальных статей (дольше)")
    ft = sub.add_parser("fulltext"); ft.add_argument("query", nargs="?", default="")
    ft.add_argument("--pmid"); ft.add_argument("--doi")
    ft.add_argument("--chars", type=int, default=20000)
    ft.add_argument("--all", action="store_true", help="выдать ВЕСЬ текст без обрезки")
    w = sub.add_parser("word"); w.add_argument("query", nargs="?", default="")
    w.add_argument("--pmid"); w.add_argument("--doi"); w.add_argument("--pdf")
    w.add_argument("--send", action="store_true")
    a = ap.parse_args()
    if a.cmd == "alerts":
        cmd_alerts(a); return
    if a.cmd == "fulltext":
        cmd_fulltext(a); return
    if a.cmd == "word" and a.pdf:
        make_word_from_pdf(a.pdf, send=a.send); return
    if a.cmd == "search":
        for i, art in enumerate(epmc_search(a.query, a.n), 1):
            m = fmt_meta(art)
            print(f"{i}. {m['title']}")
            print(f"   {m['journal']} {m['year']} | PMID {m['pmid']} | DOI {m['doi']} | {'OA' if m['oa'] else 'абстракт'}")
    else:
        make_word(a.query, pmid=a.pmid, doi=a.doi, send=a.send)

main()
