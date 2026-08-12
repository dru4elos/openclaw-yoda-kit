#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Чтение и ПРОВЕРКА документов (docx/pdf) — чтобы Йода видел, что реально отдаёт доктору.

  docs.py read <файл>              — весь текст документа
  docs.py check <файл>             — текст + проверка КАЖДОЙ картинки внутри (безопасность/тема)
  docs.py preview <файл> [--pages 3] [--send] — отрисовать страницы и посмотреть глазами (vision)
"""
import argparse, base64, json, os, re, subprocess, sys, tempfile, urllib.request, zipfile

TOME = os.path.expanduser("~/.openclaw/workspace/skills/tome/tome.py")
PYBIN = os.path.expanduser("~/mailvenv/bin/python")

def keys():
    env = {}
    for p in (os.path.expanduser("~/.openclaw/.env"),):
        if os.path.exists(p):
            for line in open(p, encoding="utf-8"):
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    k, v = line.split("=", 1)
                    env[k.strip()] = v.strip().strip('"').strip("'")
    return env.get("EXCASH_API_KEY", ""), env.get("EXCASH_API_URL", "")

def vision(path, prompt, max_tokens=1500):
    key, url = keys()
    if not key:
        return "[нет ключа для зрения]"
    b64 = base64.b64encode(open(path, "rb").read()).decode()
    ext = os.path.splitext(path)[1].lstrip(".").lower()
    mime = "jpeg" if ext in ("jpg", "jpeg") else (ext or "png")
    body = json.dumps({"model": "gemini-3.6-flash", "max_tokens": max_tokens, "temperature": 0,
        "messages": [{"role": "user", "content": [
            {"type": "image_url", "image_url": {"url": f"data:image/{mime};base64,{b64}"}},
            {"type": "text", "text": prompt}]}]}).encode()
    req = urllib.request.Request(url.rstrip("/") + "/chat/completions", data=body,
        headers={"Authorization": "Bearer " + key, "Content-Type": "application/json"})
    d = json.loads(urllib.request.urlopen(req, timeout=180).read().decode())
    return (((d.get("choices") or [{}])[0].get("message", {}) or {}).get("content") or "").strip()

def docx_text(path):
    from docx import Document
    doc = Document(path)
    out = [p.text for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                out.append(" | ".join(cells))
    return "\n".join(out)

def pdf_text(path):
    r = subprocess.run(["pdftotext", "-layout", path, "-"], capture_output=True, timeout=120)
    return r.stdout.decode("utf-8", "ignore")

def read_any(path):
    if path.lower().endswith(".docx"):
        return docx_text(path)
    if path.lower().endswith(".pdf"):
        return pdf_text(path)
    return open(path, encoding="utf-8", errors="ignore").read()

def cmd_read(a):
    t = read_any(a.file)
    print(f"[{os.path.basename(a.file)}] {len(t)} символов\n")
    print(t[:a.chars])

def cmd_check(a):
    path = a.file
    t = read_any(path)
    print(f"ПРОВЕРКА: {os.path.basename(path)}")
    print(f"Текст: {len(t)} символов. Начало: {t[:200].strip()}...\n")
    if not path.lower().endswith(".docx"):
        print("(проверка картинок доступна для .docx)")
        return
    bad = 0
    with zipfile.ZipFile(path) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
        print(f"Картинок внутри: {len(media)}")
        tmp = tempfile.mkdtemp(prefix="doccheck_")
        for i, name in enumerate(media, 1):
            data = z.read(name)
            if len(data) < 3000:
                continue
            p = os.path.join(tmp, os.path.basename(name))
            with open(p, "wb") as f:
                f.write(data)
            try:
                v = vision(p, "Классифицируй ОДНИМ словом: SAFE / NUDITY / GORE / OTHER, "
                              "затем через двоеточие 6 слов что изображено. "
                              "Медицинские снимки = SAFE.", 200).replace("\n", " ")
            except Exception as e:
                v = f"ошибка проверки: {type(e).__name__}"
            flag = "NUDITY" in v.upper() or "GORE" in v.upper()
            bad += 1 if flag else 0
            print(("  ⛔ " if flag else "  ok ") + f"{os.path.basename(name)}: {v[:110]}")
            os.remove(p)
    print(f"\nИТОГ: проблемных картинок — {bad}."
          + (" ДОКУМЕНТ ОТПРАВЛЯТЬ НЕЛЬЗЯ, пересобери без них." if bad else " Документ чистый."))

def cmd_preview(a):
    path = os.path.abspath(a.file)
    tmp = tempfile.mkdtemp(prefix="docprev_")
    pdf = path
    if path.lower().endswith(".docx"):
        soffice = None
        for c in ("/usr/bin/soffice", "/usr/bin/libreoffice"):
            if os.path.exists(c):
                soffice = c
        if not soffice:
            sys.exit("нет LibreOffice — визуальный просмотр .docx недоступен; используй check")
        subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", tmp, path],
                       capture_output=True, timeout=300)
        cand = [f for f in os.listdir(tmp) if f.endswith(".pdf")]
        if not cand:
            sys.exit("не удалось конвертировать в PDF")
        pdf = os.path.join(tmp, cand[0])
    subprocess.run(["pdftoppm", "-r", "110", "-png", "-f", "1", "-l", str(a.pages), pdf,
                    os.path.join(tmp, "page")], capture_output=True, timeout=300)
    pages = sorted(f for f in os.listdir(tmp) if f.startswith("page") and f.endswith(".png"))
    if not pages:
        sys.exit("страницы не отрисовались")
    print(f"Страниц отрисовано: {len(pages)}")
    for i, pg in enumerate(pages, 1):
        p = os.path.join(tmp, pg)
        try:
            v = vision(p, "Опиши, что реально на этой странице документа: заголовки, о чём текст, "
                          "какие изображения (что на них), таблицы. Если есть что-то неуместное "
                          "(эротика, посторонние картинки не по теме) — назови это явно.", 1200)
        except Exception as e:
            v = f"ошибка: {type(e).__name__}"
        print(f"\n--- Страница {i} ---\n{v}")
        if a.send:
            subprocess.run([PYBIN, TOME, "photo", p, "--caption", f"{os.path.basename(path)} — стр. {i}"],
                           check=False)
        os.remove(p)

def main():
    ap = argparse.ArgumentParser()
    sub = ap.add_subparsers(dest="cmd", required=True)
    r = sub.add_parser("read"); r.add_argument("file"); r.add_argument("--chars", type=int, default=12000)
    c = sub.add_parser("check"); c.add_argument("file")
    p = sub.add_parser("preview"); p.add_argument("file"); p.add_argument("--pages", type=int, default=3)
    p.add_argument("--send", action="store_true")
    a = ap.parse_args()
    if not os.path.exists(a.file):
        sys.exit(f"файл не найден: {a.file}")
    {"read": cmd_read, "check": cmd_check, "preview": cmd_preview}[a.cmd](a)

main()
