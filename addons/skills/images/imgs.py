#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""GATE_APPLIED. Картинки из интернета для Йоды: найти, прислать доктору, вставить в Word.

  imgs.py find "kung pao chicken" --n 3 [--send]     — найти и (опц.) прислать доктору
  imgs.py url "https://..." [--send]                  — скачать конкретную картинку
  imgs.py docx --title "Заголовок" --text-file /путь/текст.md \
               --img "запрос1" --img "запрос2" [--send]  — Word-документ с текстом и картинками

Источники: Wikimedia Commons + Openverse (открытые лицензии, без ключей).
Скачанное лежит в /tmp/yoda_img_* и подчищается автоочисткой.
"""
import argparse, base64, json, os, re, subprocess, sys, tempfile, urllib.parse, urllib.request
import requests

TOME = os.path.expanduser("~/.openclaw/workspace/skills/tome/tome.py")
PYBIN = os.path.expanduser("~/mailvenv/bin/python")
UA = {"User-Agent": "YodaAssistant/1.0 (+mailto:you@example.com)"}
COMMONS = "https://commons.wikimedia.org/w/api.php"
OPENVERSE = "https://api.openverse.org/v1/images/"

# Стоки: превью помечены водяным знаком — для публикации не годятся.
STOCK_DOMAINS = (
    # витрины
    "alamy.", "shutterstock.", "gettyimages.", "istockphoto.", "stock.adobe.",
    "dreamstime.", "123rf.", "depositphotos.", "vectorstock.", "canstockphoto.",
    "agefotostock.", "photostock", "lori.ru", "stockfresh.", "bigstockphoto.",
    "fotolia.", "picfair.", "zoonar.", "imago-images.", "profimedia.",
    "freepik.", "envato.", "elements.envato", "adobestock", "pond5.",
    "eyeem.", "westend61.", "masterfile.", "superstock.", "shotshop.",
    # CDN, по которым узнаётся сток (домен витрины в URL не встречается)
    "ftcdn.net", "media.gettyimages", "media.istockphoto", "image.shutterstock",
    "thumbs.dreamstime", "previews.123rf", "st.depositphotos", "st2.depositphotos",
    "st3.depositphotos", "st4.depositphotos", "c8.alamy", "l450v.alamy",
    "images.pond5", "as1.ftcdn", "as2.ftcdn",
)

# Маркеры платного превью в пути или заголовке.
STOCK_MARKERS = ("/premium-photo/", "/premium-vector/", "premium photo",
                 "premium vector", "watermark", "royalty-free image",
                 "стоковое фото", "стоковое изображение")


def is_stock(item):
    """Похоже ли на сток с водяным знаком — по URL, источнику или заголовку."""
    url = (item.get("url") or "").lower()
    src = (item.get("src") or "").lower()
    title = (item.get("title") or "").lower()
    if any(d in url or d in src for d in STOCK_DOMAINS):
        return True
    return any(m in url or m in title for m in STOCK_MARKERS)

def search_web(q, n):
    """Реальные фото из веба (DuckDuckGo Images). Основной источник — примеры «как оно выглядит»."""
    out = []
    try:
        from ddgs import DDGS
        for x in DDGS().images(q, max_results=max(n * 3, 12)):
            u = x.get("image") or x.get("thumbnail")
            if not u:
                continue
            out.append({"url": u,
                        "title": (x.get("title") or q)[:120],
                        "src": (x.get("source") or urllib.parse.urlparse(u).netloc),
                        "lic": ""})
    except Exception as e:
        sys.stderr.write(f"web-поиск: {type(e).__name__}: {str(e)[:100]}\n")
    return out[:max(n * 3, 12)]


def search_commons(q, n):
    out = []
    try:
        r = requests.get(COMMONS, params={
            "action": "query", "format": "json", "generator": "search",
            "gsrsearch": f"filetype:bitmap {q}", "gsrlimit": str(n * 2), "gsrnamespace": "6",
            "prop": "imageinfo", "iiprop": "url|extmetadata", "iiurlwidth": "1200"}, headers=UA, timeout=40)
        r.raise_for_status()
        pages = (r.json().get("query") or {}).get("pages") or {}
        for p in pages.values():
            ii = (p.get("imageinfo") or [{}])[0]
            url = ii.get("thumburl") or ii.get("url")
            if not url:
                continue
            meta = ii.get("extmetadata") or {}
            lic = (meta.get("LicenseShortName") or {}).get("value", "")
            title = re.sub(r"^File:", "", p.get("title", ""))
            out.append({"url": url, "title": title, "src": "Wikimedia Commons", "lic": lic})
    except Exception as e:
        sys.stderr.write(f"commons: {e}\n")
    return out[:n]

def search_openverse(q, n):
    out = []
    try:
        r = requests.get(OPENVERSE, params={"q": q, "page_size": str(n)}, headers=UA, timeout=40)
        if r.status_code != 200:
            return []
        for it in (r.json().get("results") or []):
            url = it.get("url")
            if url:
                out.append({"url": url, "title": it.get("title") or q,
                            "src": it.get("source") or "Openverse", "lic": it.get("license") or ""})
    except Exception as e:
        sys.stderr.write(f"openverse: {e}\n")
    return out[:n]

def download(url, idx=0):
    try:
        r = requests.get(url, headers=UA, timeout=60, stream=True)
        r.raise_for_status()
        ct = (r.headers.get("Content-Type") or "").lower()
        ext = ".jpg"
        for e in ("jpeg", "png", "webp", "gif"):
            if e in ct:
                ext = ".jpg" if e == "jpeg" else "." + e
        data = r.content[:15 * 1024 * 1024]
        if len(data) < 3000:
            return None
        p = f"/tmp/yoda_img_{idx}_{abs(hash(url)) % 10000}{ext}"
        with open(p, "wb") as f:
            f.write(data)
        os.chmod(p, 0o644)
        return p
    except Exception as e:
        sys.stderr.write(f"download {url[:60]}: {e}\n")
        return None

def find_images(q, n, allow_stock=False):
    """Веб -> Wikimedia -> Openverse, кандидатов с запасом (часть отсеет зрение).

    Стоки с водяными знаками уходят в конец списка и помечаются: публиковать их
    нельзя. Если чистых картинок нет вовсе — они всё же вернутся последними,
    но с явной пометкой, чтобы Йода предупредил доктора.
    """
    res = search_web(q, n * 2)
    res += search_commons(q, n)
    if len(res) < n * 3:
        res += search_openverse(q, n)

    seen, clean, stock = set(), [], []
    for it in res:
        u = it.get("url")
        if not u or u in seen:
            continue
        seen.add(u)
        if is_stock(it):
            it["lic"] = (it.get("lic") or "") + " ⚠️СТОК/водяной знак — публиковать нельзя"
            stock.append(it)
        else:
            clean.append(it)
    return clean + ([] if (clean and not allow_stock) else stock)

def _html_via_browser(page_url):
    """Рендер страницы настоящим Chrome — для сайтов, где картинки подгружает JS."""
    r = subprocess.run(["/opt/chrome-headless/chrome-headless-shell", "--headless",
                        "--disable-gpu", "--no-sandbox", "--virtual-time-budget=9000",
                        "--dump-dom", page_url],
                       capture_output=True, timeout=180)
    return r.stdout.decode("utf-8", "ignore")


def images_from_page(page_url, n, use_browser=False):
    """Достаёт картинки с конкретной страницы (меню, карточка отеля/блюда и т.п.)."""
    from bs4 import BeautifulSoup
    from urllib.parse import urljoin
    final_url = page_url
    if use_browser:
        html = _html_via_browser(page_url)
        if len(html) < 500:
            raise RuntimeError("браузер вернул пустую страницу")
        class _R:
            url = page_url
        r = _R()
    else:
        r = requests.get(page_url, headers={
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 "
                          "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml"}, timeout=60)
        r.raise_for_status()
        html = r.text
    soup = BeautifulSoup(html, "html.parser")
    cands, seen = [], set()

    def add(u, title=""):
        if not u:
            return
        u = urljoin(str(r.url), u.strip())
        if not u.startswith("http") or u in seen:
            return
        low = u.lower().split("?")[0]
        if low.endswith((".svg", ".ico")):
            return
        if any(w in low for w in ("logo", "icon", "sprite", "favicon", "pixel", "avatar", "banner-ad")):
            return
        seen.add(u)
        cands.append({"url": u, "title": title[:100] or "изображение со страницы",
                      "src": urllib.parse.urlparse(u).netloc, "lic": ""})

    og = soup.find("meta", property="og:image")
    if og:
        add(og.get("content"), "превью страницы")
    for img in soup.find_all("img"):
        src = (img.get("src") or img.get("data-src") or img.get("data-original")
               or img.get("data-lazy-src") or "")
        ss = img.get("srcset") or img.get("data-srcset") or ""
        if ss:
            parts = [p.strip().split(" ")[0] for p in ss.split(",") if p.strip()]
            if parts:
                src = parts[-1]
        add(src, img.get("alt") or img.get("title") or "")
    for tag in soup.find_all(style=True):
        m = re.search(r"url\((['\"]?)(.+?)\1\)", tag.get("style", ""))
        if m:
            add(m.group(2), "фон")
    return cands[:max(n * 4, 12)]


def cmd_page(a):
    try:
        cands = images_from_page(a.link, a.n, use_browser=getattr(a, "browser", False))
    except Exception as e:
        sys.exit(f"не удалось открыть страницу: {type(e).__name__}: {str(e)[:120]}")
    if not cands:
        sys.exit("на странице не нашлось картинок (возможно, они подгружаются скриптом)")
    print(f"кандидатов на странице: {len(cands)}, беру подходящие до {a.n}")
    got = 0
    for i, it in enumerate(cands, 1):
        if got >= a.n:
            break
        p = download(it["url"], i)
        if not p:
            continue
        if os.path.getsize(p) < 25000:   # мелочь: иконки/кнопки
            os.remove(p); continue
        ok, label, why = gate(p, getattr(a, "topic", "") or "")
        if not ok:
            print(f"   ⛔ отклонено фильтром [{label}]: {why}")
            os.remove(p); continue
        got += 1
        print(f"{got}. {it['title'][:60]} | {it['src']}")
        if a.send:
            subprocess.run([PYBIN, TOME, "photo", p, "--caption", it["title"][:120]], check=False)
            os.remove(p)
            print("   ✓ отправлено доктору")
        else:
            print(f"   {p}")
    if not got:
        sys.exit("подходящих картинок не нашлось (все слишком мелкие)")



def _excash_keys():
    env = {}
    for p in (os.path.expanduser("~/.openclaw/.env"), "/home/knee_bot/keys.env"):
        if not os.path.exists(p):
            continue
        try:
            fh = open(p, encoding="utf-8")
        except Exception:
            continue
        with fh:
            for line in fh:
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    k, v = line.split("=", 1)
                    env.setdefault(k.strip(), v.strip().strip(chr(34)).strip(chr(39)))
    return env.get("EXCASH_API_KEY", ""), env.get("EXCASH_API_URL", "")


def gate(path, topic=""):
    """Проверка картинки зрением. Возвращает (ok, метка, описание).
    Не прошла проверку или сомнение -> НЕ используем."""
    key, url = _excash_keys()
    if not key:
        return False, "NOKEY", "нет ключа для проверки — картинку не используем"
    try:
        b64 = base64.b64encode(open(path, "rb").read()).decode()
        ext = os.path.splitext(path)[1].lstrip(".").lower()
        mime = "jpeg" if ext in ("jpg", "jpeg") else (ext or "jpeg")
        prompt = (
            "Ты фильтр безопасности для документов врача. Оцени изображение.\n"
            "Ответь СТРОГО в формате: МЕТКА | РЕЛЕВАНТНО | краткое описание\n"
            "МЕТКА: SAFE (обычное фото места/еды/предмета/схемы/документа), "
            "NUDITY (обнажённое тело, бельё, эротика, порнография), "
            "GORE (кровь, травмы, операционное поле), OTHER (непонятно).\n"
            "РЕЛЕВАНТНО: YES или NO — подходит ли изображение теме: " + (topic or "не указана") + "\n"
            "Медицинские снимки (рентген, МРТ, КТ) считай SAFE."
        )
        body = json.dumps({"model": "gemini-3.6-flash", "max_tokens": 300, "temperature": 0,
            "messages": [{"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": f"data:image/{mime};base64,{b64}"}},
                {"type": "text", "text": prompt}]}]}).encode()
        req = urllib.request.Request(url.rstrip("/") + "/chat/completions", data=body,
            headers={"Authorization": "Bearer " + key, "Content-Type": "application/json"})
        d = json.loads(urllib.request.urlopen(req, timeout=120).read().decode())
        out = (((d.get("choices") or [{}])[0].get("message", {}) or {}).get("content") or "").strip()
    except Exception as e:
        return False, "ERR", f"проверка не удалась ({type(e).__name__}) — картинку не используем"
    up = out.upper()
    if "NUDITY" in up or "GORE" in up:
        return False, "BLOCK", out[:150]
    # Разбираем строго по полям «МЕТКА | РЕЛЕВАНТНО | описание».
    fields = [f.strip() for f in up.split("|")]
    label = fields[0] if fields else ""
    relev = fields[1] if len(fields) > 1 else ""
    if topic:
        if relev.startswith("NO") or relev == "N":
            return False, "OFFTOPIC", out[:150]
        if not relev.startswith("YES"):          # непонятный вердикт = не рискуем
            return False, "UNSURE", out[:150]
    if "SAFE" not in label:
        return False, "UNSURE", out[:150]
    return True, "SAFE", out[:150]


def cmd_find(a):
    res = find_images(a.query, a.n, allow_stock=getattr(a, "allow_stock", False))
    if not res:
        sys.exit(f"по запросу «{a.query}» ничего не нашлось")
    n_stock = sum(1 for it in res if is_stock(it))
    print(f"кандидатов: {len(res)} (из них стоков с водяным знаком: {n_stock})")
    rejected = []
    passed = 0
    for i, it in enumerate(res, 1):
        if passed >= a.n:
            break
        print(f"{i}. {it['title'][:70]} | {it['src']} {it['lic']}")
        p = download(it["url"], i) if a.send else None
        if a.send and p:
            ok, label, why = gate(p, a.query)
            if not ok:
                rejected.append(f"{it['title'][:40]} [{label}]")
                print(f"   ⛔ отклонено фильтром [{label}]: {why}")
                try:
                    os.remove(p)
                except Exception:
                    pass
                continue
            cap = f"{it['title'][:120]} ({it['src']})"
            subprocess.run([PYBIN, TOME, "photo", p, "--caption", cap], check=False)
            try:
                os.remove(p)
            except Exception:
                pass
            passed += 1
            print(f"   ✓ отправлено доктору")
        elif not a.send:
            passed += 1
            print(f"   {it['url']}")
    print(f"\nИТОГ ПРОВЕРКИ: прошло {passed} из запрошенных {a.n}; "
          f"отклонено зрением: {len(rejected)}")
    if rejected:
        print("Отклонены: " + "; ".join(rejected[:6]))
    if passed < a.n:
        print("⚠️ Набрано меньше, чем просили. Переформулируй запрос "
              "(другие слова, английский) и повтори — не выдавай неполный результат молча.")


def cmd_url(a):
    p = download(a.link, 1)
    if not p:
        sys.exit("не удалось скачать картинку")
    ok, label, why = gate(p, a.caption or "")
    if not ok:
        os.remove(p)
        sys.exit(f"⛔ картинка отклонена фильтром [{label}]: {why}")
    if a.send:
        subprocess.run([PYBIN, TOME, "photo", p, "--caption", a.caption or ""], check=False)
        os.remove(p)
        print("✓ отправлено доктору")
    else:
        print(p)

def cmd_docx(a):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    body = ""
    if a.text_file and os.path.exists(a.text_file):
        body = open(a.text_file, encoding="utf-8").read()
    elif a.text:
        body = a.text
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(12)
    doc.add_heading(a.title, level=0)

    def add_bold(p, text):
        for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
            run = p.add_run(chunk)
            if i % 2 == 1:
                run.bold = True

    for raw in body.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if re.match(r"^#{1,3}\s", line):
            doc.add_heading(line.lstrip("# ").strip(), level=min(len(line) - len(line.lstrip("#")), 3))
        elif re.match(r"^[-*•]\s", line):
            add_bold(doc.add_paragraph(style="List Bullet"), re.sub(r"^[-*•]\s+", "", line))
        elif re.match(r"^\d+[.)]\s", line):
            add_bold(doc.add_paragraph(style="List Number"), re.sub(r"^\d+[.)]\s+", "", line))
        else:
            add_bold(doc.add_paragraph(), line)

    tmp = []
    for qi, q in enumerate(a.img or [], 1):
        inserted = False
        for it in find_images(q, 4):
            if inserted:
                break
            p = download(it["url"], qi)
            if not p:
                continue
            ok, label, why = gate(p, q)
            if not ok:
                sys.stderr.write(f"пропущена картинка по «{q}» [{label}]: {why}\n")
                try:
                    os.remove(p)
                except Exception:
                    pass
                continue
            try:
                try:
                    doc.add_picture(p, width=Inches(5.5))
                except Exception:
                    from PIL import Image as _I
                    conv = p + ".conv.jpg"
                    _I.open(p).convert("RGB").save(conv, quality=90)
                    doc.add_picture(conv, width=Inches(5.5))
                    tmp.append(conv)
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph(f"{q} — {it['title'][:80]} ({it['src']}, {it['lic']})")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].italic = True
                tmp.append(p)
                inserted = True
            except Exception as e:
                sys.stderr.write(f"вставка {q}: {e}\n")
    safe = re.sub(r"[^\w\-]", "_", a.title)[:40] or "doc"
    out = f"/tmp/yoda_{safe}.docx"
    doc.save(out)
    for p in tmp:
        try:
            os.remove(p)
        except Exception:
            pass
    if a.send:
        subprocess.run([PYBIN, TOME, "file", out, "--caption", a.title[:120]], check=False)
        os.remove(out)
        print(f"✓ документ отправлен доктору ({len(tmp)} иллюстраций)")
    else:
        print(out)

def main():
    ap = argparse.ArgumentParser()
    sub = ap.add_subparsers(dest="cmd", required=True)
    f = sub.add_parser("find"); f.add_argument("query"); f.add_argument("--n", type=int, default=3)
    f.add_argument("--send", action="store_true")
    f.add_argument("--allow-stock", action="store_true",
                   help="разрешить стоки с водяными знаками (по умолчанию отсеиваются)")
    u = sub.add_parser("url"); u.add_argument("link"); u.add_argument("--caption", default="")
    u.add_argument("--send", action="store_true")
    pg = sub.add_parser("page"); pg.add_argument("link"); pg.add_argument("--n", type=int, default=5)
    pg.add_argument("--send", action="store_true")
    pg.add_argument("--topic", default="",
                    help="тема: картинки не по теме отбрасываются зрением")
    pg.add_argument("--browser", action="store_true",
                    help="рендерить страницу настоящим Chrome (JS-сайты)")
    d = sub.add_parser("docx"); d.add_argument("--title", required=True)
    d.add_argument("--text-file"); d.add_argument("--text", default="")
    d.add_argument("--img", action="append"); d.add_argument("--send", action="store_true")
    a = ap.parse_args()
    {"find": cmd_find, "url": cmd_url, "page": cmd_page, "docx": cmd_docx}[a.cmd](a)


if __name__ == "__main__":
    main()
